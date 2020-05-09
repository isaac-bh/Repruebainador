import argparse
import cv2
import docx
import imutils
import numpy as np
import os
import pathlib
import PIL.Image
import platform
import pytesseract
import sys
import time
import os
from docx import Document
from docx.shared import Pt
from imutils import contours
from imutils.perspective import four_point_transform
from io import open
from os import remove
from pytesseract import image_to_string
from PIL import Image



ruta_carpeta = sys.argv[1]
plantilla = sys.argv[2]

def Non_Zero(ruta_carpeta,plantilla):
    homedir = os.path.expanduser("~")
    directorio_documentos = homedir + "\\Documents\\Repruebainador\\Words\\"
    Path = pathlib.Path(directorio_documentos)
    for imagen in Path.iterdir():
        remove(imagen)

    Path = pathlib.Path(ruta_carpeta)
    for imagen in Path.iterdir():
        compr_ext = comprobacion_extension(str(imagen))
        if compr_ext == True:
            img = Image.open(imagen)
            new_img = img.resize((2734,3500))
            new_img.save('ajuste.png','png')

            img = cv2.imread('ajuste.png')
            crop_img = img[780:850, 1590:1710]
            cv2.imwrite('aver.png', crop_img)

            so = platform.system()
            if so == "Windows":
                pytesseract.pytesseract.tesseract_cmd = 'C:/Program Files/Tesseract-OCR/tesseract'
                TESSDATA_PREFIX = 'C:/Program Files/Tesseract-OCR'
            output = pytesseract.image_to_string(PIL.Image.open('aver.png').convert("RGB"), lang='eng', config='--psm 4 --oem 3')

            if output == 'Dias:':

                
                img = cv2.imread('ajuste.png')
                crop_img = img[850:3350, 270:1410]
                cv2.imwrite('recorte.png', crop_img)
                primero = 10
                segundo = 80
                contador = 1
                output = 'uno'
                while output != "":
                    
                    contador1 = str(contador)

                    img = cv2.imread('recorte.png')
                    crop_img = img[primero:segundo, 10:220]
                    cv2.imwrite('codigo.png', crop_img)

                    img = cv2.imread('recorte.png')
                    crop_img = img[primero:segundo, 220:1080]
                    cv2.imwrite('nombre.png', crop_img)

                    output = pytesseract.image_to_string(PIL.Image.open('codigo.png').convert("RGB"), lang='eng', config='--psm 4 --oem 3')
                    salida = pytesseract.image_to_string(PIL.Image.open('nombre.png').convert("RGB"), lang='eng', config='--psm 4 --oem 3')
                    if salida != '':
                        doc = docx.Document(plantilla) 
                        for paragraph in doc.paragraphs: 
                            if 'Nombre' in paragraph.text: 
                                paragraph.text = salida
                                run = paragraph.runs[0]
                                font = run.font
                                font.name = 'Arial'
                                font.size = Pt(12)

                            if 'Código' in paragraph.text: 
                                paragraph.text = output
                                run = paragraph.runs[0]
                                font = run.font
                                font.name = 'Arial'
                                font.size = Pt(12)

                        doc.save(directorio_documentos + salida + '.docx') 


                    primero = primero + 73
                    segundo = segundo + 73
                    contador = contador + 1

            else:

                
                img = cv2.imread('ajuste.png')
                crop_img = img[790:3290, 270:1410]
                cv2.imwrite('recorte.png', crop_img)
                
                primero = 10
                segundo = 80
                contador = 1
                output = 'uno'
                while output != "":
                    
                    contador1 = str(contador)

                    img = cv2.imread('recorte.png')
                    crop_img = img[primero:segundo, 10:220]
                    cv2.imwrite('codigo.png', crop_img)

                    img = cv2.imread('recorte.png')
                    crop_img = img[primero:segundo, 220:1080]
                    cv2.imwrite('nombre.png', crop_img)

                    output = pytesseract.image_to_string(PIL.Image.open('codigo.png').convert("RGB"), lang='eng', config='--psm 4 --oem 3')
                    salida = pytesseract.image_to_string(PIL.Image.open('nombre.png').convert("RGB"), lang='eng', config='--psm 4 --oem 3')
                    if salida != '':
                        doc = docx.Document(plantilla) 
                        for paragraph in doc.paragraphs:
                            if 'Nombre' in paragraph.text: 
                                paragraph.text = salida
                                run = paragraph.runs[0]
                                font = run.font
                                font.name = 'Arial'
                                font.size = Pt(12)

                            if 'Código' in paragraph.text: 
                                paragraph.text = output
                                run = paragraph.runs[0]
                                font = run.font
                                font.name = 'Arial'
                                font.size = Pt(12)

                        doc.save(directorio_documentos + salida + '.docx') 

                    primero = primero + 73
                    segundo = segundo + 73
                    contador = contador + 1
        else:
            pass


def eliminar_residuales():
    remove('ajuste.png')
    remove('aver.png')
    remove('codigo.png')
    remove('nombre.png')
    remove('recorte.png')

def comprobacion_extension(parametroX):
    ext = parametroX.split(".")
    if ext[1] == "png" or ext[1] == "jpg" or ext[1] == "jpeg":
        return True
    else:
        return False

Non_Zero(ruta_carpeta,plantilla)
eliminar_residuales()